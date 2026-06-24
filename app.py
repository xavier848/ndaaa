"""
HEMA NDA generator - web service (Flask).

This wraps the NDA generator so it can be called over HTTP by Power Automate.
It does the same thing as the script: takes intake data as JSON, fills the
master template, handles conditional clauses, returns the finished .docx.

Endpoints:
  GET  /            -> health check ("ok"), so you can confirm it's running in a browser
  POST /generate    -> body = intake JSON; returns the generated .docx (base64 in JSON)

Security: requires header  x-api-key  matching the SHARED_SECRET environment variable.

Run locally:   python app.py
Deployed:      a host (Render/Railway/etc.) runs it and gives an HTTPS URL.
"""

import os
import io
import re
import json
import base64

from flask import Flask, request, jsonify
from docx import Document
from docx.oxml.ns import qn

app = Flask(__name__)

MASTER_PATH = os.path.join(os.path.dirname(__file__), "HEMA_NDA_master.docx")
SHARED_SECRET = os.environ.get("NDA_SHARED_SECRET", "")


def _normalize(intake):
    """Map the form's human-readable dropdown values to the exact keywords/phrases
    the template's [[IF]] conditions and {{tokens}} expect. The form stores long
    labels (e.g. 'Standard (until ...)'); the template checks for 'Standard'.
    This is the single place that bridges that gap, so the form can keep nice labels."""
    intake = dict(intake or {})

    def low(k):
        return str(intake.get(k, "")).strip().lower()

    # DurationChoice -> "Standard" / "Fixed term"
    dc = low("DurationChoice")
    if dc.startswith("standard"):
        intake["DurationChoice"] = "Standard"
    elif dc.startswith("fixed"):
        intake["DurationChoice"] = "Fixed term"
    else:
        intake["DurationChoice"] = "TBD"   # blank/unknown -> visible placeholder, never an empty heading

    # FranchiseCarveOut -> "Keep" / "Remove"  (defensive; already clean today)
    fc = low("FranchiseCarveOut")
    if fc.startswith("keep"):
        intake["FranchiseCarveOut"] = "Keep"
    elif fc.startswith("remove"):
        intake["FranchiseCarveOut"] = "Remove"

    # NdaType -> "Mutual" / "One-sided"  (NOTE: not used by the template yet)
    nt = low("NdaType")
    if "mutual" in nt or "wederkerig" in nt:
        intake["NdaType"] = "Mutual"
    elif "one" in nt or "eenzijdig" in nt:
        intake["NdaType"] = "One-sided"

    # CounterpartyKind -> "Individual" / "Company"  (drives the parties-block variant)
    ct = low("CounterpartyType")
    is_individual = ct in ("individual", "person", "persoon", "natuurlijk persoon", "een natuurlijk persoon")
    intake["CounterpartyKind"] = "Individual" if is_individual else "Company"

    # CounterpartyType -> Dutch legal-entity phrase used in the company parties sentence
    if is_individual:
        intake["CounterpartyType"] = "een natuurlijk persoon"
    elif ct:
        intake["CounterpartyType"] = "een besloten vennootschap"

    # For an individual, "Functie" in the signature block doesn't apply -> sign in person
    if is_individual and low("CounterpartyFunction") in ("", "individual", "person", "persoon"):
        intake["CounterpartyFunction"] = "namens zichzelf"

    # CounterpartyDesignation -> how the counterparty is named throughout.
    # One-sided supplier relationship -> "Leverancier"; mutual -> neutral "Wederpartij".
    intake["CounterpartyDesignation"] = "Wederpartij" if intake.get("NdaType") == "Mutual" else "Leverancier"

    return intake



def _tokens(intake):
    return {
        "HemaName": intake.get("HemaName", "HEMA B.V."),
        "HemaKvK": (intake.get("HemaKvK") or "________"),
        "HemaSignatory": (intake.get("HemaSignatory") or "____________________"),
        "HemaFunction": intake.get("HemaFunction", "unitmanager"),
        "CounterpartyName": (intake.get("CounterpartyName") or "____________________"),
        "CounterpartyType": intake.get("CounterpartyType", "een besloten vennootschap"),
        "CounterpartyOffice": (intake.get("CounterpartyOffice") or "____________________"),
        "CounterpartyKvK": (intake.get("CounterpartyKvK") or "________"),
        "CounterpartySignatory": (intake.get("CounterpartySignatory") or "____________________"),
        "CounterpartyFunction": (intake.get("CounterpartyFunction") or "____________________"),
        "Purpose": (intake.get("Purpose") or "____________________"),
        "DurationTerm": (intake.get("DurationTerm") or "____________________"),
        "CounterpartyDesignation": intake.get("CounterpartyDesignation", "Leverancier"),
        "VABOffice": (intake.get("VABOffice") or "____________________"),
        "VABKvK": (intake.get("VABKvK") or "________"),
    }


def _cond_ok(field, value, intake):
    return str(intake.get(field, "")).strip().lower() == value.strip().lower()


def _process_text(text, intake, tok):
    def if_repl(m):
        field, value, body = m.group(1), m.group(2), m.group(3)
        return body if _cond_ok(field, value, intake) else ""
    text = re.sub(r"\[\[IF (\w+)=([^\]]+)\]\](.*?)\[\[END\]\]", if_repl, text, flags=re.S)
    text = re.sub(r"\{\{(\w+)\}\}", lambda m: str(tok.get(m.group(1), m.group(0))), text)
    return text


def _set_text(p, new):
    if p.runs:
        first = p.runs[0]
        for r in p.runs[1:]:
            r.text = ""
        first.text = new
    elif new:
        p.add_run(new)


def _delete_paragraph(p):
    p._element.getparent().remove(p._element)


def _accept_tracked_changes(doc):
    body = doc.element.body
    for ins in list(body.iter(qn('w:ins'))):
        for child in list(ins):
            ins.addprevious(child)
        ins.getparent().remove(ins)
    for dele in list(body.iter(qn('w:del'))):
        dele.getparent().remove(dele)


def generate_docx_bytes(intake):
    intake = _normalize(intake)
    doc = Document(MASTER_PATH)
    tok = _tokens(intake)

    def process(paragraphs):
        for p in list(paragraphs):
            raw = p.text
            if "[[REMOVE]]" in raw:
                _delete_paragraph(p)
                continue
            if "[[" in raw or "{{" in raw:
                new = _process_text(raw, intake, tok)
                if new.strip() == "":
                    _delete_paragraph(p)
                else:
                    _set_text(p, new)

    process(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process(cell.paragraphs)
    _accept_tracked_changes(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route("/", methods=["GET"])
def health():
    return "HEMA NDA generator is running.", 200


@app.route("/generate", methods=["POST"])
def generate():
    # auth
    if not SHARED_SECRET or request.headers.get("x-api-key") != SHARED_SECRET:
        return jsonify({"error": "unauthorized"}), 401
    # parse
    try:
        intake = request.get_json(force=True)
    except Exception as e:
        return jsonify({"error": "bad request", "detail": str(e)}), 400
    # generate
    try:
        docx_bytes = generate_docx_bytes(intake)
    except Exception as e:
        return jsonify({"error": "generation failed", "detail": str(e)}), 500
    name = (intake.get("CounterpartyName") or "NDA").replace(" ", "_")
    return jsonify({
        "filename": f"NDA_{name}.docx",
        "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "documentBase64": base64.b64encode(docx_bytes).decode("utf-8"),
    }), 200


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
